from docx import Document
import os

documents = {}

for document in os.listdir('documents'):
    documents[f'd{os.listdir("documents").index(document)+1}'] = Document(f'./documents/{document}')

new = Document()

for document in documents.values():
    paragraphs = document.paragraphs
    for para in paragraphs:
        new_para = new.add_paragraph(para.text)
        new_para.style = para.style
        new_para.alignment = para.alignment
new.save("documents/new.docx")

# for document in os.listdir('documents'):
#     if '.docx' != document[-5:]:
#         continue
#     with open(f"documents/{document}", 'rb') as file:
#         content = file.read()
#     with open("documents/new.docx", 'ab') as output:
#         output.write(content)
